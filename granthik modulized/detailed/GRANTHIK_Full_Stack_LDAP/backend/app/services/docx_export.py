"""
Service for exporting summaries to DOCX format
"""
import os
import io
import logging
from typing import Optional, List, Dict, Any
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

logger = logging.getLogger("uvicorn")

def create_summary_docx(
    title: str,
    summary_text: str,
    metadata: Optional[Dict[str, Any]] = None
) -> bytes:
    """
    Create a DOCX document with the summary
    
    Args:
        title: Title for the document
        summary_text: Summary text to include
        metadata: Optional metadata to include
        
    Returns:
        Bytes of the DOCX document
    """
    try:
        # Create document
        doc = Document()
        
        # Set document properties
        doc.core_properties.title = title
        doc.core_properties.author = "GRANTHIK"
        
        # Add title
        title_para = doc.add_heading(title, level=1)
        title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Add metadata if provided
        if metadata:
            metadata_para = doc.add_paragraph()
            metadata_para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            
            for key, value in metadata.items():
                if value:
                    meta_run = metadata_para.add_run(f"{key}: {value}")
                    meta_run.italic = True
                    meta_run.font.size = Pt(10)
                    metadata_para.add_run("\n")
        
        # Add divider
        doc.add_paragraph("_" * 50)
        
        # Add summary text
        summary_para = doc.add_paragraph()
        summary_para.add_run(summary_text)
        
        # Save document to bytes
        buffer = io.BytesIO()
        doc.save(buffer)
        buffer.seek(0)
        
        return buffer.getvalue()
    
    except Exception as e:
        logger.error(f"Error creating DOCX summary: {str(e)}")
        import traceback
        logger.error(traceback.format_exc())
        raise