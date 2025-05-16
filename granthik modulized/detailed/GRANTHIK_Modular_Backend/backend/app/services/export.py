"""
Enhanced export service for exporting summaries to different formats with markdown support
"""
import os
import io
import re
import logging
from datetime import datetime
from typing import Dict, Any, Optional, List, Tuple

from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

# For PDF export
from reportlab.lib.pagesizes import letter
from reportlab.lib import colors
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle

# For HTML export
import markdown
from weasyprint import HTML

logger = logging.getLogger("uvicorn")

def parse_markdown(text: str) -> List[Tuple[str, str, str]]:
    """
    Parse markdown text into a list of (type, content, extra) tuples
    
    Args:
        text: Markdown text to parse
        
    Returns:
        List of (type, content, extra) tuples where:
        - type: heading1, heading2, heading3, paragraph, bullet, numbered, etc.
        - content: The text content
        - extra: Additional info like indentation level
    """
    lines = text.split('\n')
    parsed = []
    
    for line in lines:
        line = line.rstrip()
        
        # Skip empty lines
        if not line:
            parsed.append(('empty', '', ''))
            continue
        
        # Headings
        if line.startswith('# '):
            parsed.append(('heading1', line[2:], ''))
        elif line.startswith('## '):
            parsed.append(('heading2', line[3:], ''))
        elif line.startswith('### '):
            parsed.append(('heading3', line[4:], ''))
        elif line.startswith('#### '):
            parsed.append(('heading4', line[5:], ''))
        
        # Lists
        elif line.startswith('- ') or line.startswith('* '):
            indent = len(line) - len(line.lstrip())
            level = indent // 2  # Estimate indentation level
            parsed.append(('bullet', line[2:].strip(), str(level)))
        elif re.match(r'^\d+\.\s', line):
            indent = len(line) - len(line.lstrip())
            level = indent // 2  # Estimate indentation level
            content = re.sub(r'^\d+\.\s', '', line).strip()
            parsed.append(('numbered', content, str(level)))
        
        # Code blocks
        elif line.startswith('```'):
            parsed.append(('codeblock', line[3:], 'start'))
        elif line.endswith('```'):
            parsed.append(('codeblock', line[:-3], 'end'))
        
        # Regular paragraph
        else:
            parsed.append(('paragraph', line, ''))
    
    return parsed

def export_to_docx(
    summary_text: str, 
    filename: str, 
    model_used: Optional[str] = None,
    created_at: Optional[datetime] = None,
    summary_type: Optional[str] = None
) -> bytes:
    """
    Export summary to DOCX format with markdown support
    
    Args:
        summary_text: The summary text to export (markdown format)
        filename: Original document filename
        model_used: The model used to generate the summary
        created_at: When the summary was created
        summary_type: Type of summary (comprehensive, concise, key_points)
        
    Returns:
        Bytes of the DOCX file
    """
    try:
        # Create a new Document
        doc = Document()
        
        # Set document properties
        doc.core_properties.title = f"Summary of {filename}"
        doc.core_properties.author = "GRANTHIK"
        
        # Add styles for different heading levels
        styles = doc.styles
        
        # Add a title
        title = doc.add_heading(f"Summary of {filename}", level=1)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Add metadata
        metadata_para = doc.add_paragraph()
        if model_used:
            model_run = metadata_para.add_run(f"Generated with: {model_used}")
            model_run.italic = True
            model_run.font.size = Pt(10)
            metadata_para.add_run("\n")
        
        if created_at:
            date_run = metadata_para.add_run(f"Date: {created_at.strftime('%Y-%m-%d %H:%M:%S')}")
            date_run.italic = True
            date_run.font.size = Pt(10)
            metadata_para.add_run("\n")
            
        if summary_type:
            type_run = metadata_para.add_run(f"Summary type: {summary_type}")
            type_run.italic = True
            type_run.font.size = Pt(10)
        
        metadata_para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        
        # Add a horizontal line
        doc.add_paragraph("_" * 50)
        
        # Parse markdown and add content
        parsed_content = parse_markdown(summary_text)
        
        in_code_block = False
        for item_type, content, extra in parsed_content:
            if item_type == 'empty':
                doc.add_paragraph()
            elif item_type == 'heading1':
                doc.add_heading(content, level=1)
            elif item_type == 'heading2':
                doc.add_heading(content, level=2)
            elif item_type == 'heading3':
                doc.add_heading(content, level=3)
            elif item_type == 'heading4':
                doc.add_heading(content, level=4)
            elif item_type == 'bullet':
                level = int(extra) if extra.isdigit() else 0
                p = doc.add_paragraph(style='List Bullet')
                p.paragraph_format.left_indent = Inches(level * 0.25)
                p.add_run(content)
            elif item_type == 'numbered':
                level = int(extra) if extra.isdigit() else 0
                p = doc.add_paragraph(style='List Number')
                p.paragraph_format.left_indent = Inches(level * 0.25)
                p.add_run(content)
            elif item_type == 'codeblock':
                if extra == 'start':
                    in_code_block = True
                    p = doc.add_paragraph()
                    p.style = 'No Spacing'
                    p.paragraph_format.left_indent = Inches(0.5)
                    p.paragraph_format.right_indent = Inches(0.5)
                    run = p.add_run(content)
                    run.font.name = 'Courier New'
                    run.font.size = Pt(9)
                elif extra == 'end':
                    in_code_block = False
                    p = doc.add_paragraph()
                    p.style = 'No Spacing'
                    p.paragraph_format.left_indent = Inches(0.5)
                    p.paragraph_format.right_indent = Inches(0.5)
                    run = p.add_run(content)
                    run.font.name = 'Courier New'
                    run.font.size = Pt(9)
            elif item_type == 'paragraph':
                # Handle bold and italic formatting
                p = doc.add_paragraph()
                
                # Process bold and italic markdown
                parts = re.split(r'(\*\*.*?\*\*|\*.*?\*|`.*?`)', content)
                for part in parts:
                    if part.startswith('**') and part.endswith('**'):
                        # Bold text
                        run = p.add_run(part[2:-2])
                        run.bold = True
                    elif part.startswith('*') and part.endswith('*'):
                        # Italic text
                        run = p.add_run(part[1:-1])
                        run.italic = True
                    elif part.startswith('`') and part.endswith('`'):
                        # Code text
                        run = p.add_run(part[1:-1])
                        run.font.name = 'Courier New'
                        run.font.size = Pt(9)
                    elif part:
                        # Regular text
                        p.add_run(part)
        
        # Save the document to a bytes buffer
        buffer = io.BytesIO()
        doc.save(buffer)
        buffer.seek(0)
        
        return buffer.getvalue()
    except Exception as e:
        logger.error(f"Error exporting to DOCX: {str(e)}")
        raise

def export_to_pdf(
    summary_text: str, 
    filename: str, 
    model_used: Optional[str] = None,
    created_at: Optional[datetime] = None,
    summary_type: Optional[str] = None
) -> bytes:
    """
    Export summary to PDF format with markdown support
    
    Args:
        summary_text: The summary text to export (markdown format)
        filename: Original document filename
        model_used: The model used to generate the summary
        created_at: When the summary was created
        summary_type: Type of summary (comprehensive, concise, key_points)
        
    Returns:
        Bytes of the PDF file
    """
    try:
        # Convert markdown to HTML
        html_content = markdown.markdown(summary_text)
        
        # Create a complete HTML document
        complete_html = f"""
        <!DOCTYPE html>
        <html>
        <head>
            <meta charset="UTF-8">
            <title>Summary of {filename}</title>
            <style>
                body {{ font-family: Arial, sans-serif; margin: 40px; }}
                h1 {{ text-align: center; color: #333; }}
                .metadata {{ text-align: right; font-style: italic; color: #666; margin-bottom: 20px; }}
                hr {{ border: 1px solid #ddd; margin: 20px 0; }}
                pre {{ background-color: #f5f5f5; padding: 10px; border-radius: 5px; }}
                code {{ font-family: 'Courier New', monospace; background-color: #f5f5f5; padding: 2px 4px; }}
                ul, ol {{ margin-left: 20px; }}
            </style>
        </head>
        <body>
            <h1>Summary of {filename}</h1>
            <div class="metadata">
                {f"Generated with: {model_used}<br>" if model_used else ""}
                {f"Date: {created_at.strftime('%Y-%m-%d %H:%M:%S')}<br>" if created_at else ""}
                {f"Summary type: {summary_type}" if summary_type else ""}
            </div>
            <hr>
            {html_content}
        </body>
        </html>
        """
        
        # Convert HTML to PDF using WeasyPrint
        buffer = io.BytesIO()
        HTML(string=complete_html).write_pdf(buffer)
        buffer.seek(0)
        
        return buffer.getvalue()
    except Exception as e:
        logger.error(f"Error exporting to PDF: {str(e)}")
        raise

def export_to_txt(
    summary_text: str, 
    filename: str, 
    model_used: Optional[str] = None,
    created_at: Optional[datetime] = None,
    summary_type: Optional[str] = None
) -> bytes:
    """
    Export summary to plain text format
    
    Args:
        summary_text: The summary text to export
        filename: Original document filename
        model_used: The model used to generate the summary
        created_at: When the summary was created
        summary_type: Type of summary (comprehensive, concise, key_points)
        
    Returns:
        Bytes of the text file
    """
    try:
        # Create header
        header = f"Summary of {filename}\n"
        header += "=" * len(header) + "\n\n"
        
        # Add metadata
        metadata = ""
        if model_used:
            metadata += f"Generated with: {model_used}\n"
        if created_at:
            metadata += f"Date: {created_at.strftime('%Y-%m-%d %H:%M:%S')}\n"
        if summary_type:
            metadata += f"Summary type: {summary_type}\n"
        
        if metadata:
            metadata += "\n" + "-" * 50 + "\n\n"
        
        # Combine all parts
        full_text = header + metadata + summary_text
        
        # Convert to bytes
        return full_text.encode('utf-8')
    except Exception as e:
        logger.error(f"Error exporting to TXT: {str(e)}")
        raise

def export_to_html(
    summary_text: str, 
    filename: str, 
    model_used: Optional[str] = None,
    created_at: Optional[datetime] = None,
    summary_type: Optional[str] = None
) -> bytes:
    """
    Export summary to HTML format with markdown support
    
    Args:
        summary_text: The summary text to export (markdown format)
        filename: Original document filename
        model_used: The model used to generate the summary
        created_at: When the summary was created
        summary_type: Type of summary (comprehensive, concise, key_points)
        
    Returns:
        Bytes of the HTML file
    """
    try:
        # Convert markdown to HTML
        html_content = markdown.markdown(summary_text)
        
        # Create a complete HTML document
        complete_html = f"""
        <!DOCTYPE html>
        <html>
        <head>
            <meta charset="UTF-8">
            <title>Summary of {filename}</title>
            <style>
                body {{ font-family: Arial, sans-serif; margin: 40px; }}
                h1 {{ text-align: center; color: #333; }}
                .metadata {{ text-align: right; font-style: italic; color: #666; margin-bottom: 20px; }}
                hr {{ border: 1px solid #ddd; margin: 20px 0; }}
                pre {{ background-color: #f5f5f5; padding: 10px; border-radius: 5px; }}
                code {{ font-family: 'Courier New', monospace; background-color: #f5f5f5; padding: 2px 4px; }}
                ul, ol {{ margin-left: 20px; }}
                @media print {{
                    body {{ margin: 0.5in; }}
                    .no-print {{ display: none; }}
                }}
            </style>
        </head>
        <body>
            <h1>Summary of {filename}</h1>
            <div class="metadata">
                {f"Generated with: {model_used}<br>" if model_used else ""}
                {f"Date: {created_at.strftime('%Y-%m-%d %H:%M:%S')}<br>" if created_at else ""}
                {f"Summary type: {summary_type}" if summary_type else ""}
            </div>
            <hr>
            {html_content}
            <div class="no-print" style="text-align: center; margin-top: 30px;">
                <button onclick="window.print()">Print this summary</button>
            </div>
        </body>
        </html>
        """
        
        return complete_html.encode('utf-8')
    except Exception as e:
        logger.error(f"Error exporting to HTML: {str(e)}")
        raise