"""
Utility for exporting chat summaries to DOCX format
"""
import os
import logging
from typing import List, Dict, Any, Optional
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

logger = logging.getLogger("uvicorn")

def create_docx_summary(
    title: str,
    summary_text: str,
    chat_history: Optional[List[Dict[str, Any]]] = None,
    output_path: Optional[str] = None
) -> str:
    """
    Create a DOCX document with the summary and optionally chat history
    
    Args:
        title: Title for the document
        summary_text: Summary text to include
        chat_history: Optional list of chat messages
        output_path: Optional path to save the document
        
    Returns:
        Path to the saved document
    """
    try:
        logger.info(f"Creating DOCX summary with title: {title}")
        
        # Create document
        doc = Document()
        
        # Set document properties
        doc.core_properties.title = title
        doc.core_properties.author = "GRANTHIK"
        
        # Add title
        title_para = doc.add_paragraph()
        title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        title_run = title_para.add_run(title)
        title_run.bold = True
        title_run.font.size = Pt(16)
        
        # Add summary heading
        doc.add_heading("Summary", level=1)
        
        # Add summary text
        summary_para = doc.add_paragraph()
        summary_para.add_run(summary_text)
        
        # Add chat history if provided
        if chat_history and len(chat_history) > 0:
            doc.add_heading("Chat History", level=1)
            
            for message in chat_history:
                # Add user query
                if "query" in message and message["query"]:
                    query_para = doc.add_paragraph()
                    query_run = query_para.add_run(f"User: {message['query']}")
                    query_run.bold = True
                    query_run.font.color.rgb = RGBColor(0, 0, 128)  # Dark blue
                
                # Add AI response
                if "answer" in message and message["answer"]:
                    answer_para = doc.add_paragraph()
                    answer_para.add_run(f"AI: {message['answer']}")
                
                # Add separator
                doc.add_paragraph("---")
        
        # Save document
        if not output_path:
            # Create a temporary file path
            import tempfile
            output_dir = tempfile.gettempdir()
            safe_title = "".join(c for c in title if c.isalnum() or c in " _-").strip()
            safe_title = safe_title.replace(" ", "_")
            output_path = os.path.join(output_dir, f"{safe_title}_{os.urandom(4).hex()}.docx")
        
        doc.save(output_path)
        logger.info(f"DOCX summary saved to: {output_path}")
        
        return output_path
    
    except Exception as e:
        logger.error(f"Error creating DOCX summary: {str(e)}")
        import traceback
        logger.error(traceback.format_exc())
        raise